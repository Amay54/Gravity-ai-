import os
import uuid
from datetime import datetime

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from backend.reporting.chart_generator import ChartGenerator
from backend.reporting.citation_engine import CitationEngine
from backend.schemas.research import ResearchReport, SharedResearchContext


class DOCXGenerator:
    """
    Generates editable corporate Microsoft Word documents using python-docx.
    """

    @staticmethod
    def _get_theme_rgb(theme: str) -> tuple[RGBColor, RGBColor]:
        theme = theme.lower()
        if theme == "dark":
            return RGBColor(56, 189, 248), RGBColor(168, 85, 247)
        elif theme == "minimal":
            return RGBColor(55, 65, 81), RGBColor(107, 114, 128)
        elif theme == "corporate":
            return RGBColor(30, 58, 138), RGBColor(71, 85, 105)
        else:  # professional
            return RGBColor(79, 70, 229), RGBColor(124, 58, 237)

    @classmethod
    def generate(
        cls,
        context: SharedResearchContext,
        report_data: ResearchReport,
        session_id: str,
        version: int,
        theme: str = "Professional",
        user_name: str = "Developer",
    ) -> str:
        doc_dir = "backend/storage/docx"
        os.makedirs(doc_dir, exist_ok=True)
        docx_path = os.path.join(doc_dir, f"{session_id}_v{version}.docx")

        doc = Document()

        # Configure page margins
        for section in doc.sections:
            section.top_margin = Inches(1.0)
            section.bottom_margin = Inches(1.0)
            section.left_margin = Inches(1.0)
            section.right_margin = Inches(1.0)

            # Setup headers/footers
            footer = section.footer
            f_p = footer.paragraphs[0]
            f_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            f_p.text = f"GravityAI Corporate Intelligence Dossier | Version v{version} | Session: {session_id} | Page "

        primary_rgb, secondary_rgb = cls._get_theme_rgb(theme)

        # Title slide / Cover Page in document
        title_p = doc.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        title_run = title_p.add_run("GRAVITYAI")
        title_run.font.name = "Helvetica"
        title_run.font.size = Pt(14)
        title_run.font.bold = True
        title_run.font.color.rgb = primary_rgb

        comp_title_p = doc.add_paragraph()
        comp_run = comp_title_p.add_run(
            f"Corporate Research Dossier:\n{report_data.company_profile.name.value}"
        )
        comp_run.font.name = "Helvetica"
        comp_run.font.size = Pt(26)
        comp_run.font.bold = True
        comp_run.font.color.rgb = primary_rgb

        sub_p = doc.add_paragraph()
        sub_run = sub_p.add_run("Autonomously compiled via LangGraph Multi-Agent Research System.")
        sub_run.font.name = "Helvetica"
        sub_run.font.size = Pt(12)
        sub_run.font.color.rgb = secondary_rgb

        doc.add_paragraph("\n" * 4)

        # Cover Metadata Table
        table = doc.add_table(rows=4, cols=2)
        table.autofit = True

        meta_items = [
            ("Session ID", session_id),
            ("Date compiled", datetime.utcnow().strftime("%Y-%m-%d %H:%M UTC")),
            ("Dossier Version", f"v{version}"),
            ("Authenticated User", user_name),
        ]

        for idx, (label, val) in enumerate(meta_items):
            row_cells = table.rows[idx].cells
            row_cells[0].paragraphs[0].add_run(label).bold = True
            row_cells[1].paragraphs[0].add_run(val)

        doc.add_page_break()

        # Generate bibliography references
        references = CitationEngine.generate_references(context)

        # Helper to add section headers
        def add_heading_1(text: str):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(8)
            run = p.add_run(text)
            run.font.name = "Helvetica"
            run.font.size = Pt(16)
            run.font.bold = True
            run.font.color.rgb = primary_rgb

        def add_heading_2(text: str):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(text)
            run.font.name = "Helvetica"
            run.font.size = Pt(12)
            run.font.bold = True
            run.font.color.rgb = secondary_rgb

        def add_body_p(text: str, italic: bool = False):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(text)
            run.font.name = "Helvetica"
            run.font.size = Pt(10)
            run.font.italic = italic
            return p

        def add_factual_p(label: str, fact_obj: dict):
            val = fact_obj.get("value", "Not Available")
            if isinstance(val, list):
                val = ", ".join(val) if val else "Not Available"

            conf = fact_obj.get("confidence", 0.0)

            # Map citations
            citation_indexes = []
            evidences = fact_obj.get("evidence", [])
            for ev in evidences:
                for idx, r in enumerate(references):
                    if r["url"] == ev.get("url"):
                        citation_indexes.append(str(idx + 1))
            citations_str = (
                f" [{', '.join(sorted(set(citation_indexes)))}]" if citation_indexes else ""
            )

            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)

            run_lbl = p.add_run(f"{label}: ")
            run_lbl.bold = True
            run_lbl.font.size = Pt(10)

            run_val = p.add_run(f"{val}{citations_str}")
            run_val.font.size = Pt(10)

            # Subtitle Source/Confidence
            p_meta = doc.add_paragraph()
            p_meta.paragraph_format.space_after = Pt(8)
            run_m = p_meta.add_run(
                f"Source: {fact_obj.get('source', 'N/A')} | Confidence: {conf * 100:.0f}%"
            )
            run_m.font.size = Pt(8)
            run_m.font.italic = True
            run_m.font.color.rgb = RGBColor(107, 114, 128)

        # 1. EXECUTIVE SUMMARY
        add_heading_1("1. Executive Summary")
        add_body_p(report_data.company_profile.description.value)
        add_body_p(
            f"Research Quality Score: {report_data.metadata.research_quality_score * 100:.0f}% | Overall Confidence: {report_data.metadata.overall_confidence * 100:.0f}%"
        )

        # 2. COMPANY PROFILE
        add_heading_1("2. Company Profile & Sitemap")
        add_factual_p("HQ Location", report_data.company_profile.hq_location.model_dump())
        add_factual_p("Founded Year", report_data.company_profile.founded_year.model_dump())
        add_factual_p("Key Leadership", report_data.company_profile.key_leadership.model_dump())
        add_factual_p("Industry Classification", report_data.company_profile.industry.model_dump())

        # 3. FINANCIAL ANALYSIS
        add_heading_1("3. Financial Analysis & Business Model")
        add_heading_2("Financial Metrics")
        add_factual_p("Valuation", report_data.financial_analysis.valuation.model_dump())
        add_factual_p("Revenue Trends", report_data.financial_analysis.revenue_trends.model_dump())
        add_factual_p("Funding Stage", report_data.financial_analysis.funding_rounds.model_dump())

        add_heading_2("Business Model")
        add_factual_p(
            "Pricing Model",
            report_data.financial_analysis.business_model.pricing_model.model_dump(),
        )
        add_factual_p(
            "Revenue Streams",
            report_data.financial_analysis.business_model.revenue_streams.model_dump(),
        )
        add_factual_p(
            "Customer Segments",
            report_data.financial_analysis.business_model.customer_segments.model_dump(),
        )

        # Embedded Revenue Chart
        temp_images = []
        if report_data.financial_analysis.revenue_chart_data:
            chart_labels = report_data.financial_analysis.revenue_chart_data.get("labels", [])
            chart_values = report_data.financial_analysis.revenue_chart_data.get("data", [])
            if chart_labels and chart_values:
                _, chart_bytes = ChartGenerator.generate_revenue_chart(
                    chart_labels, chart_values, theme
                )

                temp_dir = "backend/storage/docx/tmp"
                os.makedirs(temp_dir, exist_ok=True)
                temp_filename = f"{uuid.uuid4()}.png"
                temp_path = os.path.join(temp_dir, temp_filename)
                with open(temp_path, "wb") as f:
                    f.write(chart_bytes)
                temp_images.append(temp_path)

                doc.add_picture(temp_path, width=Inches(4.5))

        # 4. TECH STACK
        add_heading_1("4. Technology Stack Scan")
        add_factual_p(
            "Frontend UI Frameworks", report_data.tech_stack.frontend_frameworks.model_dump()
        )
        add_factual_p("Backend tech / Servers", report_data.tech_stack.backend_tech.model_dump())
        add_factual_p("Databases / Storage", report_data.tech_stack.databases.model_dump())
        add_factual_p("Cloud Providers", report_data.tech_stack.cloud_providers.model_dump())
        add_factual_p("CDNs", report_data.tech_stack.cdns.model_dump())
        add_factual_p(
            "Analytics & tracking", report_data.tech_stack.analytics_platforms.model_dump()
        )

        # 5. HIRING TRENDS
        add_heading_1("5. Hiring Activity & Distribution")
        add_factual_p("Hiring Velocity", report_data.hiring_trends.hiring_velocity.model_dump())
        add_factual_p("Active vacancies count", report_data.hiring_trends.open_roles.model_dump())
        add_factual_p("Growing Departments", report_data.hiring_trends.top_departments.model_dump())

        if report_data.hiring_trends.hiring_chart_data:
            h_labels = report_data.hiring_trends.hiring_chart_data.get("labels", [])
            h_values = report_data.hiring_trends.hiring_chart_data.get("data", [])
            if h_labels and h_values:
                _, chart_bytes = ChartGenerator.generate_hiring_chart(h_labels, h_values, theme)

                temp_filename = f"{uuid.uuid4()}.png"
                temp_path = os.path.join(temp_dir, temp_filename)
                with open(temp_path, "wb") as f:
                    f.write(chart_bytes)
                temp_images.append(temp_path)

                doc.add_picture(temp_path, width=Inches(4.5))

        # 6. COMPETITORS
        add_heading_1("6. Competitor Landscape Matrix")
        add_factual_p(
            "Market Positioning Overview",
            report_data.competitor_analysis.market_positioning.model_dump(),
        )

        # Table
        comp_table = doc.add_table(rows=1, cols=3)
        comp_table.style = "Table Grid"
        hdr_cells = comp_table.rows[0].cells
        hdr_cells[0].paragraphs[0].add_run("Competitor").bold = True
        hdr_cells[1].paragraphs[0].add_run("Market Share").bold = True
        hdr_cells[2].paragraphs[0].add_run("Key Advantages").bold = True

        for comp in report_data.competitor_analysis.direct_competitors:
            row_cells = comp_table.add_row().cells
            name = comp.name if hasattr(comp, "name") else comp.get("name", "N/A")
            focus = comp.focus if hasattr(comp, "focus") else comp.get("focus", "N/A")
            comparison = (
                comp.comparison if hasattr(comp, "comparison") else comp.get("comparison", "N/A")
            )
            row_cells[0].paragraphs[0].add_run(name)
            row_cells[1].paragraphs[0].add_run(focus)
            row_cells[2].paragraphs[0].add_run(comparison)

        # 7. PATENT ACTIVITY
        add_heading_1("7. Patent Activity & Innovation")
        add_factual_p("Patent Count", report_data.patent_activity.patent_counts.model_dump())
        add_factual_p("Filing Trends", report_data.patent_activity.filing_trends.model_dump())
        add_factual_p(
            "Focus Areas", report_data.patent_activity.technology_focus_areas.model_dump()
        )

        # 8. SOCIAL
        add_heading_1("8. Official Digital Presence")
        add_factual_p(
            "LinkedIn profile", report_data.digital_presence.linkedin_profile.model_dump()
        )
        add_factual_p("GitHub Organization", report_data.digital_presence.github_org.model_dump())
        add_factual_p("YouTube channel", report_data.digital_presence.youtube_channel.model_dump())
        add_factual_p(
            "Developer portal docs", report_data.digital_presence.developer_docs.model_dump()
        )

        # 9. SWOT
        add_heading_1("9. SWOT Strategic Evaluation")
        swot = report_data.swot_matrix
        add_heading_2("Strengths")
        for s in swot.strengths:
            add_body_p(f"• {s}")
        add_heading_2("Weaknesses")
        for w in swot.weaknesses:
            add_body_p(f"• {w}")
        add_heading_2("Opportunities")
        for o in swot.opportunities:
            add_body_p(f"• {o}")
        add_heading_2("Threats")
        for t in swot.threats:
            add_body_p(f"• {t}")

        # 10. RECOMMENDATIONS
        add_heading_1("10. Strategic Recommendations")
        for idx, rec in enumerate(report_data.strategic_recommendations):
            add_body_p(f"{idx + 1}. {rec}")

        # 11. REFERENCES
        add_heading_1("11. Verbatim Bibliography & References")
        if not references:
            add_body_p("No external sitemaps or references cached.")
        else:
            for r in references:
                add_body_p(f"{r['citation_text']}")
                add_body_p(
                    f'Section: {r["section"]}.{r["field_name"]} | Verbatim Quote: "{r["quote"]}"',
                    italic=True,
                )

        doc.save(docx_path)

        # Clean up temp images
        for temp_img in temp_images:
            try:
                os.remove(temp_img)
            except Exception:
                pass

        return docx_path
